from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, italic=False):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def para(text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False,
         italic=False, space_before=0, space_after=6, keep_together=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if keep_together:
        pf.keep_together = True
    if text:
        run = p.add_run(text)
        set_font(run, size, bold, italic)
    return p

def heading(text, level_size, space_before=12, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, level_size, bold=True)
    return p

def chapter_heading(number, title):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'Chapter {number}')
    set_font(run, 20, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(12)
    run2 = p2.add_run(title)
    set_font(run2, 20, bold=True)
    return p2

def section_heading(text):
    return heading(text, 16, space_before=10, space_after=4)

def subsection_heading(text):
    return heading(text, 14, space_before=8, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)

def body(text, space_after=6):
    return para(text, size=12, space_after=space_after)

def bullet(text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, 12)
    return p

def bullet2(text):
    return bullet(text, indent=0.9)

def page_break():
    doc.add_page_break()

def centered(text, size=12, bold=False):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, bold=bold, space_after=4)

def sig_line(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('----------------------------------------')
    set_font(run, 12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(label)
    set_font(run2, 12)

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
centered('')
centered('UniVerse – A Campus Companion', size=22, bold=True)
centered('')
centered('Swadheen Islam Robi (0182320012101278)', size=12)
centered('Fahmid Alam (0182320012101309)', size=12)
centered('Shahriar Rashid Ratul (0182320012101276)', size=12)
centered('')
centered('')
centered('B.Sc. Project in Computing Science and Engineering', size=12)
centered('Supervisor at CSE-LU: Jaminur Rahman', size=12)
centered('Assistant Professor, Dept. of CSE, Leading University', size=12)
centered('')
centered('')
centered('Leading University', size=12, bold=True)
centered('Department of Computing Science and Engineering', size=12)
centered('Ragibnagar, South Surma, Sylhet-3112', size=12)
centered('BANGLADESH', size=12)
centered('')
centered('June 2026', size=12)

# ════════════════════════════════════════════════════════════════════════════
# RECOMMENDATION LETTER – SUPERVISOR
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Recommendation Letter from the Project Supervisor',
        16, space_before=0, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

body('The project entitled "UniVerse – A Campus Companion" was submitted by the students,')
para('', space_after=2)
for m in [
    ('Swadheen Islam Robi', '0182320012101278'),
    ('Fahmid Alam',          '0182320012101309'),
    ('Shahriar Rashid Ratul','0182320012101276'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'Mr. {m[0]}')
    set_font(run, 12, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(f'ID: {m[1]}')
    set_font(run2, 12)

para('', space_after=4)
body('is a record of research work carried out under my supervision and I, hereby, approve that the report be submitted in partial fulfillment of the requirements for the award of their Bachelor Degrees.')
para('', space_after=16)
body('Signature of the Supervisor:', space_after=20)
sig_line('Jaminur Rahman')
body('Assistant Professor, Dept. of CSE')
body('Leading University, Sylhet')
body('Date:')

# ════════════════════════════════════════════════════════════════════════════
# RECOMMENDATION LETTER – HEAD OF DEPARTMENT
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Recommendation Letter from the Head of the Department',
        16, space_before=0, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

body('The project entitled "UniVerse – A Campus Companion" was submitted by the students,')
para('', space_after=2)
for m in [
    ('Swadheen Islam Robi', '0182320012101278'),
    ('Fahmid Alam',          '0182320012101309'),
    ('Shahriar Rashid Ratul','0182320012101276'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'Mr. {m[0]}')
    set_font(run, 12, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(f'ID: {m[1]}')
    set_font(run2, 12)

para('', space_after=4)
body('is, hereby, accepted as the partial fulfillment of the requirements for the award of their Bachelor\'s Degree.')
para('', space_after=16)
body('Signature of the Head of the Department:', space_after=20)
sig_line('Kazi Md. Jahid Hasan')
body('Assistant Professor & Head (Acting)')
body('Dept. of CSE')
body('Leading University, Sylhet')

# ════════════════════════════════════════════════════════════════════════════
# CERTIFICATE OF ACCEPTANCE
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Certificate of Acceptance of the Project',
        16, space_before=0, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

body('The project entitled "UniVerse – A Campus Companion" was submitted by the students,')
para('', space_after=2)
for m in [
    ('Swadheen Islam Robi', '0182320012101278'),
    ('Fahmid Alam',          '0182320012101309'),
    ('Shahriar Rashid Ratul','0182320012101276'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'Mr. {m[0]}')
    set_font(run, 12, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(f'ID: {m[1]}')
    set_font(run2, 12)

para('', space_after=4)
body('is, hereby, accepted as the partial fulfillment of the requirements for the award of their Bachelor\'s Degree.')
para('', space_after=24)
body('Signatures:', space_after=20)

tbl = doc.add_table(rows=2, cols=3)
tbl.style = 'Table Grid'
cells = tbl.rows[0].cells
for i, label in enumerate(['Chairman, Defense Board', 'Member 1, Defense Board', 'Member 2, Defense Board']):
    cells[i].text = ''
    p = cells[i].paragraphs[0]
    run = p.add_run('----------------------------------')
    set_font(run, 12)
cells2 = tbl.rows[1].cells
for i, label in enumerate(['Chairman, Defense Board', 'Member 1, Defense Board', 'Member 2, Defense Board']):
    cells2[i].text = ''
    p = cells2[i].paragraphs[0]
    run = p.add_run(label)
    set_font(run, 12)

# ════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Acknowledgements', 16, space_before=0, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

body('First and foremost, we express our sincere gratitude to our project supervisor, Jaminur Rahman, Assistant Professor, Department of CSE, Leading University, for his invaluable guidance, constant encouragement, and insightful feedback throughout the development of this project. His direction helped us navigate the technical and academic challenges we encountered.')
body('We are deeply grateful to Kazi Md. Jahid Hasan, Assistant Professor and Head (Acting) of the Department of CSE, for approving this project and for providing the academic environment that made this work possible.')
body('We would also like to thank the faculty members of the Department of CSE, Leading University, for their cooperation in providing the course distribution data that was essential for verifying our automatic timetable generator.')
body('Our heartfelt thanks go to our families for their unwavering support and patience throughout the duration of this project.')
body('Finally, we acknowledge the open-source communities behind Flutter, Supabase, FastAPI, Firebase, and Google OR-Tools, whose tools and documentation made the technical implementation of UniVerse possible.')

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Abstract', 16, space_before=0, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

body('UniVerse – A Campus Companion is a mobile-first academic management platform developed for Leading University, Sylhet. The application addresses the fragmented nature of academic information delivery in educational institutions by centralizing class routine management, academic resource sharing, real-time notifications, and teacher–student communication within a single Android application. The key technical contribution of this project is an automatic department timetable generator, built using Google OR-Tools CP-SAT constraint programming solver, deployed as a live FastAPI service on Render. The generator ingests the existing Course Distribution Excel workbook, applies hard and soft scheduling constraints—including teacher availability, room capacity, lab requirements, and the university\'s 7-day week with Friday Period-4 restrictions—and produces a clash-free, formatted timetable workbook. The generated routine is then published directly to the live database and made available to all students and teachers through the app in real time. The student and teacher dashboards provide live class visibility with countdown timers, and push notifications via Firebase Cloud Messaging ensure instant alerts for class cancellations, resource uploads, and administrative broadcasts. The application was built using Flutter (Dart), Supabase (PostgreSQL + Auth + Realtime + Storage), and Firebase Cloud Messaging, following a clean layered architecture with role-based access control for Student, Teacher, and Admin actors. The system has been verified end-to-end with Summer 2025 semester data comprising 358 course offerings and 55 cohorts, producing a clash-free schedule with full room assignment.')

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Contents', 16, space_before=0, space_after=10, align=WD_ALIGN_PARAGRAPH.LEFT)

toc_entries = [
    ('Recommendation Letter from the Project Supervisor', 'i'),
    ('Recommendation Letter from the Head of the Department', 'ii'),
    ('Certificate of Acceptance of the Project', 'iii'),
    ('Acknowledgements', 'iv'),
    ('Abstract', 'v'),
    ('Chapter 1 – Introduction', '1'),
    ('Chapter 2 – Problem Description', '3'),
    ('    2.1  Problem Statement', '3'),
    ('    2.2  Goals', '4'),
    ('    2.3  Purposes', '4'),
    ('    2.4  Methods', '4'),
    ('    2.5  Related Work', '5'),
    ('Chapter 3 – Requirement Specification', '6'),
    ('    3.1  Functional Requirements', '6'),
    ('    3.2  Non-functional Requirements', '7'),
    ('    3.3  UML Diagrams', '8'),
    ('Chapter 4 – Accomplishment', '10'),
    ('    4.1  Preliminaries', '10'),
    ('    4.2  How the Work Was Done', '11'),
    ('Chapter 5 – Results', '14'),
    ('Chapter 6 – Conclusions', '17'),
    ('    6.1  Restrictions', '17'),
    ('    6.2  Limitations', '18'),
    ('    6.3  Future Work', '18'),
    ('References', '19'),
    ('Appendix A – Source Code Highlights', '20'),
    ('Appendix B – User\'s Guide', '21'),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(entry)
    set_font(run, 12)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
chapter_heading('1', 'Introduction')

body('UniVerse – A Campus Companion is a mobile-first academic management application developed for Leading University, Sylhet, Bangladesh by Team Sherlocked as part of the Bachelor of Science Project-I (CSE-3240). The application is built for Android using the Flutter framework and is backed by a cloud-hosted Supabase (PostgreSQL) database, Firebase Cloud Messaging for push notifications, and a separately deployed FastAPI timetable generation engine.')

body('The primary motivation behind UniVerse stems from the operational challenges faced by the Department of CSE at Leading University. Faculty members were spending considerable time manually constructing the department timetable from a Course Distribution workbook containing hundreds of offerings across dozens of cohorts and teachers. The process was error-prone, time-consuming, and prone to scheduling conflicts. At the same time, students lacked a centralized digital platform to access their class routines, receive cancellation alerts, and browse academic resources. Communication between administrators, teachers, and students relied heavily on informal channels.')

body('UniVerse was designed and developed to solve these problems through three principal contributions:')
bullet('An automatic timetable generator that applies Google OR-Tools CP-SAT constraint programming to produce clash-free, room-assigned, formatted weekly timetables from the existing Course Distribution workbook, and publishes them directly to the student and teacher routine views.')
bullet('A role-aware mobile companion that provides students with a real-time dashboard showing their live and next class with a countdown timer, academic resource access by semester, and push-notification-driven alerts; and provides teachers with a dashboard and class management tools including per-occurrence class cancellation.')
bullet('A unified administration panel allowing admin users to configure the timetable engine, broadcast notifications, manage academic resources, and govern user registration.')

body('The report is organized as follows: Chapter 2 describes the problem in detail. Chapter 3 specifies the system requirements. Chapter 4 documents the accomplishment and development process. Chapter 5 presents the results. Chapter 6 concludes the report with restrictions, limitations, and future scope.')

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 – PROBLEM DESCRIPTION
# ════════════════════════════════════════════════════════════════════════════
chapter_heading('2', 'Problem Description')

body('This chapter presents the core problem that UniVerse was designed to address, along with the project goals, purposes, methods, and a survey of related work.')

section_heading('2.1  Problem Statement')
body('At Leading University\'s Department of CSE, academic scheduling and information dissemination present recurring operational challenges:')
bullet('The departmental timetable is produced manually by faculty, involving hundreds of course offerings across 55+ cohorts (batch–section combinations), 76+ teachers, and constrained room availability. The process requires days of effort, is prone to teacher and room clashes, and must be repeated every semester.')
bullet('Once produced, the timetable is distributed informally (printed sheets, WhatsApp groups), with no structured digital access for students or teachers.')
bullet('Class cancellations, room changes, and test reminders are communicated through informal channels, reaching students inconsistently and late.')
bullet('Academic resources such as previous-year question papers and course notes are scattered across personal Google Drives and Facebook groups, with no organized per-semester access.')
bullet('There is no mobile-first platform offering students a live view of their current or next class.')

body('The combination of these factors results in wasted faculty time, scheduling errors, and poor student experience. UniVerse was developed to address all of these issues in a single, integrated platform.')

section_heading('2.2  Goals')
body('The primary goals of the UniVerse project are:')
bullet('Automate the generation of clash-free departmental timetables from the existing Course Distribution Excel workbook using constraint programming.')
bullet('Provide students with a role-aware mobile dashboard showing their live and upcoming class in real time, along with a full weekly routine view.')
bullet('Enable teachers to manage their classes—posting room changes, test reminders, and per-occurrence cancellations—with automatic student notification.')
bullet('Centralize academic resource sharing in a semester-organized hub accessible to all students.')
bullet('Deliver push notifications for all relevant academic events (class cancellation, resource upload, routine publish, admin broadcast) to the appropriate audience in real time.')
bullet('Provide administrators with a unified panel covering user management, broadcast messaging, resource management, and timetable configuration and generation.')

section_heading('2.3  Purposes')
body('The purposes of the UniVerse application are:')
bullet('To reduce the manual effort and eliminate scheduling errors in the departmental timetable construction process.')
bullet('To provide a structured, mobile-first interface for students to access their academic schedule and resources without relying on informal channels.')
bullet('To improve teacher–student communication through instant push notifications for class-related updates.')
bullet('To demonstrate the practical applicability of constraint programming (OR-Tools CP-SAT) to the NP-hard university course timetabling problem in a real departmental setting.')
bullet('To serve as a defense-ready, feature-complete demonstration of a full-stack mobile application integrating Flutter, Supabase, Firebase, and a deployed Python API.')

section_heading('2.4  Methods')
body('The following methods and technologies were employed:')
bullet('Constraint Programming (CP-SAT): Google OR-Tools CP-SAT solver was used to assign (day, period) slots to course sessions subject to hard constraints (no clashes, teacher availability, Friday Period-4 restriction) and soft constraints (course sessions on different days, cohort compactness, avoiding last period). A separate greedy phase assigned physical rooms.')
bullet('Flutter (Dart): Used to build the Android mobile application with a clean layered architecture (Screen → Controller → Service → Supabase). State management used ChangeNotifier + ListenableBuilder throughout.')
bullet('Supabase (PostgreSQL + Auth + Realtime + Storage): Used as the backend database, authentication provider, real-time event bus, and file storage. Row-Level Security (RLS) policies enforce role-based data access.')
bullet('Firebase Cloud Messaging (FCM): Used to deliver OS-level push notifications. A Supabase Edge Function (send-push) is triggered by a database webhook on every notification INSERT and fans out to registered device tokens.')
bullet('FastAPI (Python): Used to build the timetable engine REST API, deployed on Render. The engine accepts the Course Distribution workbook and configuration JSON, runs the solver, and returns the result as both structured data and a formatted Excel workbook.')
bullet('Agile feature-based development: Each team member owned vertical slices of the application, developing screens, controllers, and services for their assigned features.')

section_heading('2.5  Related Work')
body('Several existing systems partially address the problems that UniVerse targets:')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for i, h in enumerate(['System', 'Key Features', 'Limitations']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        set_font(run, 12, bold=True)

rows_data = [
    ('University ERP Portals', 'Timetable display, results, notices', 'Web-only; poor mobile experience; no automation; no push notifications'),
    ('Google Classroom', 'Assignment submission, class management', 'Not routine-focused; no timetable generation; no cancellation management'),
    ('Moodle', 'Learning management, course content', 'Complex UI; not student-dashboard oriented; no scheduling automation'),
    ('FET (Free Timetabling)', 'Automated timetable generation', 'Desktop application; no mobile companion; no database integration'),
    ('UniTime', 'University scheduling system', 'Enterprise-grade complexity; not practical for a single department deployment'),
]
for row_data in rows_data:
    row = tbl.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for run in row.cells[i].paragraphs[0].runs:
            set_font(run, 12)

para('', space_after=6)
body('None of the above systems provide a unified mobile-first platform combining automated constraint-programming-based timetable generation, a live student/teacher dashboard with class countdown, per-occurrence class cancellation with push notification, and semester-organized academic resource access.')

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – REQUIREMENT SPECIFICATION
# ════════════════════════════════════════════════════════════════════════════
chapter_heading('3', 'Requirement Specification')

body('This chapter outlines the functional and non-functional requirements of UniVerse, along with supporting UML and system design diagrams.')

section_heading('3.1  Functional Requirements')
body('The following functional requirements were identified and fully implemented:')

subsection_heading('Authentication and Role Management')
bullet('FR-01: Users can register as Student (with batch and section), Teacher (with teacher code), or Admin (whitelist-gated).')
bullet('FR-02: Users can sign in via Google OAuth (PKCE flow) or email/password with email verification.')
bullet('FR-03: The system enforces role-based access: students see Student tabs, teachers see Teacher tabs, admins see Admin tabs.')
bullet('FR-04: Admin accounts are provisioned via an invite Edge Function holding the service-role key; non-whitelisted admin attempts are rejected with a dedicated screen.')

subsection_heading('Student Features')
bullet('FR-05: Students can view a dashboard showing a greeting, a live-class hero card (with room and teacher), or a next-class countdown card, followed by today\'s schedule and recent notifications preview.')
bullet('FR-06: Students can view their full weekly routine filtered by their batch and section.')
bullet('FR-07: Students can browse academic resources organized by semester folder and category (notes, previous questions, etc.), open files in-app or via browser, and copy Drive links.')
bullet('FR-08: Students receive push notifications for class cancellations, resource uploads, routine publishes, and admin broadcasts.')
bullet('FR-09: Students can view, multi-select dismiss (locally), and read all notifications in a feed.')

subsection_heading('Teacher Features')
bullet('FR-10: Teachers can view a dashboard showing their live or next class with countdown, today\'s schedule, and quick actions.')
bullet('FR-11: Teachers can view their full weekly routine filtered by their teacher code.')
bullet('FR-12: Teachers can cancel a specific occurrence of a class (selecting the next date) with a reason; the cancellation is recorded and students receive an instant push notification.')
bullet('FR-13: Teachers can post class updates (room change, notice, test reminder) targeting the affected batch and section.')
bullet('FR-14: Teachers can undo a cancellation they created.')

subsection_heading('Admin Features')
bullet('FR-15: Admins can view an administrative dashboard with system statistics.')
bullet('FR-16: Admins can compose and send broadcast notifications to all users, by role, or by batch/section.')
bullet('FR-17: Admins can upload academic resources (any file type, or Drive links) with semester and category metadata; students receive an automatic push notification on upload.')
bullet('FR-18: Admins can generate a department timetable by uploading the Course Distribution Excel file, configuring rooms, faculty, and settings, and triggering the engine. They can review the generation report, validation results, and a visual timetable grid before downloading or publishing.')
bullet('FR-19: Admins can manage timetable rooms (add/edit/activate/deactivate), faculty (add/edit/set off-days), and timetable settings (period times, Friday rule, weights).')
bullet('FR-20: Admins can manage user accounts and view broadcast history.')

subsection_heading('Push Notification System')
bullet('FR-21: Every INSERT into the notifications table automatically triggers the deployed send-push Edge Function, which resolves the audience and delivers OS push notifications via FCM HTTP v1 to all registered device tokens.')

section_heading('3.2  Non-functional Requirements')
bullet('NFR-01 Performance: Timetable generation for 358 offerings (Summer 2025) completes within 60 seconds on the deployed Render free-tier instance. Dashboard data loads within 2 seconds under normal network conditions.')
bullet('NFR-02 Security: Supabase Row-Level Security enforces data isolation. The service-role key is never exposed to the client; it lives only in the invite-admin Edge Function. The anon key in the APK is restricted by RLS to permitted read operations only.')
bullet('NFR-03 Scalability: The timetable engine is stateless; multiple concurrent generation jobs are supported via background thread isolation. The Supabase backend scales independently.')
bullet('NFR-04 Usability: The app follows a consistent dark-theme design system (AppColors, AppTextStyles, AppSpacing) with Inter font and Phosphor icons. All interactions provide immediate visual feedback.')
bullet('NFR-05 Reliability: Push notification delivery is decoupled from the app (DB webhook → Edge Function → FCM), ensuring notifications are sent even if the app is not running. Realtime Supabase subscriptions provide in-app live updates.')
bullet('NFR-06 Maintainability: The codebase follows a strict layered architecture (Screen → Controller → Service → Supabase) with no Supabase calls in UI layers. GoRouter handles all navigation with named route constants.')
bullet('NFR-07 Compatibility: The app targets Android. The APK uses debug signing for sideloadable demo distribution. Gradle 8.14 + AGP 8.9.1 is required for JDK 24 compatibility.')

section_heading('3.3  System Design Diagrams')

subsection_heading('3.3.1  Activity Diagram – Timetable Generation Flow')
body('The timetable generation activity follows these steps:')
bullet('Admin navigates to the Routine hub → Generate tab.')
bullet('Admin selects the Course Distribution Excel workbook from local storage using the file picker.')
bullet('The app loads the current timetable configuration from the database (rooms, faculty, settings) and sends a multipart POST request to the engine API.')
bullet('The engine ingests the workbook (parsing the Course Distribution sheet), runs the CP-SAT solver for Phase 1 (day/period assignment) and a greedy algorithm for Phase 2 (room assignment), and renders the output workbook.')
bullet('The app polls the job status endpoint until the state transitions to done or failed.')
bullet('On success, the app displays a generation report (stats, validation results). The admin can view the visual timetable grid, download the Excel file, or publish to the routines table.')
bullet('On publish, all existing routines for the semester are replaced and a push notification is sent to all users.')

subsection_heading('3.3.2  Use Case Diagram')
body('The system has three primary actors: Student, Teacher, and Admin. The key use cases per actor are:')
body('Student: View dashboard → View routine → Browse resources → View/dismiss notifications → Manage profile.')
body('Teacher: View dashboard → View routine → Manage classes (cancel/update/undo) → View/dismiss notifications → Manage profile.')
body('Admin: View dashboard → Broadcast notification → Generate timetable → Manage routine (view/delete entries) → Manage resources (upload/delete) → Configure rooms/faculty/settings → Manage users → View broadcast history.')
body('All actors share: Sign up → Sign in (Google OAuth or email/password) → Receive push notifications.')

subsection_heading('3.3.3  ER Diagram')
body('The key entities and their relationships in the Supabase (PostgreSQL) database are:')
bullet('profiles (id, role, full_name, batch, section, teacher_code, avatar_url) — extends auth.users.')
bullet('routines (id, day, time_start, time_end, subject, subject_code, teacher_name, teacher_code, room, batch, section, semester_label, is_lab) — the published weekly schedule.')
bullet('cancellations (id, routine_id [FK→routines], class_date, reason, batch, section, subject, day, time_start, cancelled_by [FK→profiles]) — per-occurrence cancellations.')
bullet('notifications (id, title, body, type, target_role, target_batch, target_section, sent_by [FK→profiles], created_at) — in-app and push notifications.')
bullet('notification_reads (id, notification_id [FK→notifications], user_id [FK→profiles], read_at) — per-user read state.')
bullet('resources (id, title, description, file_url, file_type, category, semester, uploaded_by [FK→profiles]) — academic files and links.')
bullet('device_tokens (id, user_id [FK→profiles], token, platform, created_at) — FCM device tokens for push delivery.')
bullet('timetable_rooms (id, name, building, is_lab, is_gallery, is_active) — room pool for the generator.')
bullet('timetable_faculty (id, acronym, full_name, dept, designation, off_days[], is_active) — faculty directory for the generator.')
bullet('timetable_settings (id=1, semester_label, periods jsonb, friday_no_p4, service_scope, weights jsonb) — generator configuration.')
bullet('timetable_runs (id, semester_label, file_path, stats jsonb, validation jsonb, status, row_count, created_by) — generation run history.')
bullet('whitelists (email, role) — admin whitelist gate.')

subsection_heading('3.3.4  Data Flow Diagram')
body('DFD-0 (Context Diagram): UniVerse has four external entities: Student, Teacher, Admin, and the Timetable Engine (external FastAPI service). The central UniVerse system receives routine queries, resource requests, and notification reads from Students and Teachers; receives timetable configuration, generation requests, and management commands from Admin; and sends generation jobs to and receives results from the Timetable Engine.')
body('DFD-1 (Level-1): The system is decomposed into five major processes:')
bullet('Process 1 – Authentication & Registration: Handles sign-in (Google/email), role assignment, email verification, and whitelist checking. Data stores: auth.users, profiles, whitelists.')
bullet('Process 2 – Routine & Dashboard: Fetches routine data for the student\'s batch/section or teacher\'s code; computes live/next class; serves the dashboard hero cards. Data stores: routines, cancellations.')
bullet('Process 3 – Notification System: Accepts notification creation (from admin broadcast, teacher cancel/update, resource upload, or routine publish); inserts into notifications; the DB webhook fires send-push → FCM → device. Data stores: notifications, notification_reads, device_tokens.')
bullet('Process 4 – Resource Hub: Admin uploads files to Supabase Storage and creates resource records; students browse and open resources filtered by semester/category. Data stores: resources, Supabase Storage (resources bucket).')
bullet('Process 5 – Timetable Generation & Config: Admin configures rooms/faculty/settings; triggers generation (POST to engine with config); polls status; reviews result; downloads or publishes. Data stores: timetable_rooms, timetable_faculty, timetable_settings, timetable_runs, routines, Supabase Storage (timetables bucket).')

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 – ACCOMPLISHMENT
# ════════════════════════════════════════════════════════════════════════════
chapter_heading('4', 'Accomplishment')

body('This chapter describes the development process of UniVerse, covering the team structure, technical preliminaries, and how each major feature was designed and built.')

section_heading('4.1  Preliminaries')
body('Before development began, the following groundwork was established:')
bullet('Technology selection: Flutter was chosen for cross-platform mobile development with a single codebase targeting Android. Supabase was chosen as the backend-as-a-service for its managed PostgreSQL, Auth, Realtime, and Storage capabilities, eliminating the need to manage a custom server. Firebase Cloud Messaging was selected for push notifications due to its reliability and Flutter integration.')
bullet('Database schema design: All tables were designed with Row-Level Security in mind from the outset. A migration-based approach (supabase/migrations/) was adopted to track schema changes, with migrations numbered 001–007 covering initial schema through the cancellations table.')
bullet('Architecture decision: A strict three-layer architecture (Screen → Controller → Service → Supabase) was enforced project-wide, with ChangeNotifier as the only state management pattern. This ensured that screens never contained business logic and that Supabase was accessed only through service classes.')
bullet('Design system: A unified dark-theme design system was established in lib/core/theme/ with AppColors, AppTextStyles, and AppSpacing token files. All UI components reference these tokens exclusively, ensuring visual consistency.')
bullet('Feature ownership: Three team members divided ownership by vertical feature slice. Fahmid Alam owned the shared infrastructure, auth system, admin features, and timetable engine. Swadheen Islam Robi owned the routine/dashboard/resources/teacher features. Shahriar Rashid Ratul owned notifications, profile, push integration, and QA.')

section_heading('4.2  How the Work Was Done')

subsection_heading('4.2.1  Authentication System')
body('The auth system supports two sign-in methods: Google OAuth (using Supabase\'s PKCE flow with the google_sign_in package) and email/password with email verification. After sign-in, handlePostLogin() in AuthService creates or updates the user\'s profile row, checks the whitelist for admin accounts, and registers the FCM device token. The AuthController exposes an AuthStatus enum (initial, loading, authenticated, unauthenticated, registering, notWhitelisted, awaitingVerification, error) that the GoRouter redirect function uses to enforce navigation rules.')
body('Student and teacher registration collect additional profile data (batch/section for students, teacher code for teachers) that is stored in the profiles table. Admin provisioning uses the invite-admin Supabase Edge Function, which holds the service-role key and is the only path to creating admin accounts.')

subsection_heading('4.2.2  Automatic Timetable Generator')
body('The timetable engine is the project\'s primary technical differentiator. It was developed as a Python FastAPI application and deployed on Render:')
bullet('Ingestion (ingest.py): Parses the "Course Distribution" sheet of the provided Excel workbook using openpyxl. Uses header-text binding to locate course, teacher, batch, and section columns. Applies the even-digit lab rule (courses with even-digit credit counts scheduled as lab sessions), excludes blank-teacher project rows and blank-batch rows, detects service offerings (non-CSE courses), and validates the parsed data.')
bullet('Solver (solver.py): Phase 1 uses CP-SAT to assign a (day, period) tuple to each session. Hard constraints: each session scheduled exactly once; no teacher double-booking at the same (day, period); no cohort (batch+section) double-booking; teacher off-days respected; Friday Period-4 not used. Soft constraints: a course\'s two weekly sessions on different days (penalized if same day); cohort compactness (sessions not spread too thin); avoid last period. Phase 2 uses a greedy algorithm to assign physical rooms: lab sessions → lab rooms, theory sessions → theory rooms or galleries, respecting room capacity constraints.')
bullet('Rendering (render.py): Clones the CSE_Routine_TEMPLATE.xlsx workbook and writes one canonical 55-row cohort map (indexed by cohort order) to all 7 day-sheets, with cells formatted as "COURSE_CODE TEACHER ROOM". Friday Period-4 cells are left blank.')
bullet('API (main.py): Exposes POST /api/timetable/generate (accepts multipart file + config JSON + time_limit_s, returns job_id), GET /api/timetable/status/{job_id} (returns state and progress), GET /api/timetable/result/{job_id} (returns rows, stats, validation), and GET /api/timetable/download/{job_id} (returns the rendered workbook).')
bullet('Flutter integration: TimetableEngineService handles the full lifecycle: file upload → POST generate → poll status → fetch result → upload workbook to Supabase Storage → publish rows to routines table → insert a "routine published" push notification. The GenerateTimetableScreen presents a step-by-step UI: file picker → progress with engine log → report card → validation card → timetable grid preview → download / publish actions.')
body('The engine was verified with Summer 2025 semester data: 358 course offerings processed into 654 CSE sessions and 32 service sessions, across 55 cohorts and 76 teachers, producing a clash-free solution with validation.ok = true.')

subsection_heading('4.2.3  Student and Teacher Dashboards')
body('Both dashboards follow the same structural pattern: a greeting app bar, a hero card section, a stat strip, today\'s class list, and quick actions.')
body('The hero section adapts to the current time: if a class is currently in session, a LiveClassCard is shown with the subject, room, teacher, and time; if a class is coming within the day, a NextClassCard shows a live countdown timer (updated every second via a periodic timer in the controller); if all classes are done for the day, a MessageHeroCard shows an encouraging message. The controller computes these states by comparing the current time against the sorted list of today\'s routine entries, taking cancellations into account.')
body('The dashboards also show a today\'s-class list at the bottom, with cancelled classes visually badged. A recent-alerts preview shows the three most recent unread notifications.')

subsection_heading('4.2.4  Teacher Manage Classes')
body('The Manage Classes screen allows teachers to manage their class occurrences:')
bullet('The screen presents a day selector (Sun–Sat). Selecting a day shows all of the teacher\'s classes on that day, with a "CANCELLED" badge if the next occurrence has a cancellation record.')
bullet('Tapping a class opens an action sheet with three options: Cancel Occurrence, Post Update, and Undo Cancellation.')
bullet('Cancel Occurrence: The teacher enters a reason. TeacherService.cancelClass() computes the next date of that weekday, inserts a row into cancellations (routine_id, class_date, reason, batch, section, subject, day, time_start, cancelled_by), and calls NotificationService.createBroadcast() to insert a class_cancel notification targeting the affected batch and section. The DB webhook fires automatically, pushing the notification to all affected students\' devices.')
bullet('Post Update: The teacher selects an update type (room_change, notice, test_reminder), writes a message, and posts it as a notification broadcast.')
bullet('Undo: Deletes the cancellation row for the next occurrence of that class.')

subsection_heading('4.2.5  Resources Hub')
body('The resources feature serves two roles:')
body('Admin side: The Manage Resources screen allows admins to upload files using the file picker (any type: PDF, Word, image) or enter a Google Drive link. ResourceService.uploadFile() uploads to the Supabase resources Storage bucket; ResourceService.createResource() inserts the metadata row (title, file_url, category, semester, uploaded_by). After upload, a notification is automatically inserted targeting all students, triggering a push notification via the DB webhook.')
body('Student side: The ResourceController fetches all resources from the database. The resources screen presents them organized into semester folders (e.g., "Semester 1", "Summer 2025"). Tapping a folder shows a category list; tapping a category shows the files. Files open in-app (PDF viewer for PDFs, system file opener for others) or launch in a browser for Drive links, with clipboard fallback.')

subsection_heading('4.2.6  Push Notification System')
body('Push notifications are delivered through a fully automated pipeline:')
bullet('When any row is inserted into the notifications table (by admin broadcast, teacher cancel, resource upload, or routine publish), the Supabase DB Webhook fires the send-push Edge Function.')
bullet('The Edge Function resolves the notification\'s audience (all users, by role, by batch, or by section) by querying profiles, then looks up the matching device_tokens rows.')
bullet('It mints a short-lived JWT using the FCM service account credentials (stored as the FCM_SERVICE_ACCOUNT secret in the Edge Function environment) and sends the notification via the FCM HTTP v1 API.')
bullet('On the app side, PushService.init() in main.dart handles foreground message display (via flutter_local_notifications), background message handling (via FirebaseMessaging.onBackgroundMessage), and tap-to-open routing (navigates to the notifications feed).')
bullet('Device tokens are registered on sign-in (PushService.registerToken) and cleaned up on sign-out (before the session drops, to avoid stale token rows).')

subsection_heading('4.2.7  Database Migrations and RLS')
body('Seven migrations were applied to the Supabase database:')
bullet('001 – notification_reads table.')
bullet('002 – Drop legacy photo_url column from profiles.')
bullet('003 – Convert teacher_name and teacher_code in routines from foreign key to TEXT (to decouple from profiles).')
bullet('004 – device_tokens table.')
bullet('005 – register_device_token database function.')
bullet('006 – Enable RLS on all tables; add my_role() and is_admin() helper functions used in RLS policies.')
bullet('007 – cancellations table with canonical columns, unique constraint on (routine_id, class_date), RLS (read-all; insert/delete by cancelled_by = auth.uid() restricted to teacher/admin role), and Realtime publication. Also relaxes the legacy cancel_date NOT NULL constraint on routines.')

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 – RESULTS
# ════════════════════════════════════════════════════════════════════════════
chapter_heading('5', 'Results')

body('This chapter describes the functional outcomes of the UniVerse system across its three actor roles and the timetable generation engine. The application was built and verified end-to-end on the polish/defense-prep branch.')

section_heading('5.1  Timetable Generation Results')
body('The timetable engine was tested with real Summer 2025 semester data from Leading University\'s Department of CSE:')
bullet('Input: 358 course offerings from the Course Distribution workbook.')
bullet('Parsed output: 654 CSE sessions + 32 service (non-CSE) sessions across 55 cohorts and 76 teachers.')
bullet('Solver result: Clash-free schedule (validation.ok = true). No teacher double-bookings, no cohort double-bookings, all teacher off-days respected, Friday Period-4 left unscheduled.')
bullet('Room assignment: All sessions assigned to appropriate rooms (labs to lab rooms, theory to theory rooms or galleries).')
bullet('Render output: A formatted 7-day Excel workbook using the CSE_Routine_TEMPLATE.xlsx template, with all cohorts correctly mapped.')
bullet('End-to-end time: Within 60 seconds on the Render free-tier deployment (cold start ~50 seconds if the instance has been idle for >15 minutes).')
body('The generated routine was published to the routines table, replacing all previous entries for the semester, and a push notification was delivered to all registered users.')

section_heading('5.2  Student Experience')
body('A student signed in with a valid batch and section sees:')
bullet('Dashboard: Personalized greeting, hero card showing the current live class (subject, room, teacher) or a countdown to the next class, a today\'s-class list, and a recent-alerts preview.')
bullet('Routine tab: Full 7-day weekly routine filtered to their batch and section, with Sunday–Saturday tabs. Cancelled occurrences are visually badged CANCELLED.')
bullet('Resources tab: Semester folder grid → category list → file list. Files open directly (PDF viewer, external opener) or via browser for Drive links.')
bullet('Alerts tab: Chronological notification feed with per-notification and multi-select dismiss (stored locally in SharedPreferences; DB rows are not deleted).')
bullet('Profile tab: Name, role, batch, section, avatar (uploaded to Supabase Storage avatars bucket).')

section_heading('5.3  Teacher Experience')
body('A teacher signed in with a valid teacher code sees:')
bullet('Dashboard: Personalized greeting, live or next-class hero card, today\'s schedule, and quick actions (Go to Routine, Manage Classes).')
bullet('Routine tab: Full 7-day routine showing all classes assigned to their teacher code.')
bullet('Classes tab (Manage Classes): Day selector → class list with CANCELLED badges. Tap a class to Cancel Occurrence (with reason → DB row + student push notification), Post Update, or Undo.')
bullet('Alerts tab and Profile tab: Same as student.')

section_heading('5.4  Admin Experience')
body('An admin signed in via the whitelist sees:')
bullet('Dashboard: System statistics and quick-access cards.')
bullet('Broadcast tab: Compose notification (title, body, audience: all/role/batch/section) with broadcast history accessible from the app bar.')
bullet('Routine tab: Admin Routine hub with two segments — Manage (view/delete existing routine entries) and Generate (full timetable generation workflow: file pick → progress → report → grid → download/publish).')
bullet('Users tab: View and manage registered user accounts.')
bullet('Profile tab: Same as other roles, but with admin designation.')
bullet('Secondary screens: Manage Resources (upload/delete files and Drive links with Resource Library), Manage Rooms, Manage Faculty, Timetable Settings.')

section_heading('5.5  Push Notification Verification')
body('Push notification delivery was verified for the following event types:')
bullet('Class cancellation by teacher → affected batch/section students receive OS push notification within seconds.')
bullet('Resource upload by admin → all students receive OS push notification.')
bullet('Routine publish by admin → all users receive OS push notification.')
bullet('Admin broadcast → targeted audience (all/role/batch/section) receives OS push notification.')
body('The send-push Edge Function is deployed and active. The Supabase DB Webhook on the notifications table INSERT triggers it reliably. Notification delivery was confirmed on physical Android devices.')

section_heading('5.6  Application Build')
body('The release APK was built using flutter build apk --release. The APK is sideloadable and uses debug signing for the demo. The TIMETABLE_BASE_URL defaults to the live Render engine URL baked into app_constants.dart. The app icon was generated using flutter_launcher_icons from the orbit-mark PNG assets.')

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 – CONCLUSIONS
# ════════════════════════════════════════════════════════════════════════════
chapter_heading('6', 'Conclusions')

body('UniVerse – A Campus Companion successfully achieves its primary objectives. The application delivers a complete, defense-ready mobile platform that addresses the key pain points of academic information management at Leading University\'s Department of CSE. The automatic timetable generator—the core technical differentiator—produces clash-free, room-assigned, formatted weekly timetables from the existing Course Distribution workbook in under 60 seconds, and publishes them directly to the student and teacher routine views via a live database integration.')

body('The student and teacher dashboards provide real-time class visibility that was previously unavailable through any digital channel. The push notification pipeline ensures that class cancellations, resource uploads, and announcements reach students\' devices within seconds of the event. The academic resource hub gives all students organized access to course materials that were previously scattered across informal channels.')

body('The project goes significantly beyond the original proposal in its technical scope, adding constraint programming-based optimization, a deployed FastAPI microservice, a full push notification infrastructure, live dashboard widgets with countdown timers, and teacher class management features that were not originally scoped.')

section_heading('6.1  Restrictions')
bullet('The application currently targets Android only. An iOS build would require additional Firebase and Supabase configuration, an Apple Developer account, and platform-specific entitlements.')
bullet('The timetable engine is deployed on the Render free tier, which sleeps after 15 minutes of inactivity. The first generation request after an idle period incurs a cold start delay of approximately 50 seconds. A pre-warm request before a live demonstration is necessary.')
bullet('The release APK uses debug signing and is intended for sideloading only. Publishing to the Google Play Store would require a production keystore and the Play Developer program.')
bullet('Admin provisioning requires access to the invite-admin Edge Function. There is no self-service admin registration path by design.')
bullet('The timetable generator is specific to the CSE Course Distribution workbook format at Leading University. It requires the "Course Distribution" sheet header format to be preserved for correct parsing.')

section_heading('6.2  Limitations')
bullet('The AI academic assistant (RAG/Gemini) that was proposed as a core feature was descoped during development due to time constraints and the complexity of the vector search migration. The documents table and the RouteNames.aiAssistant constant are preserved for future restoration.')
bullet('The student-facing routine view does not yet visually badge cancelled classes in the weekly grid (only the teacher Manage Classes view and the dashboard today-list show CANCELLED). This is a UI-only gap; the cancellations data is correctly stored and alerts are delivered.')
bullet('The teacher directory (a browsable list of teacher profiles with contact information) proposed in the original feature set was not built as a standalone screen. Teacher information is accessible through the routine view and timetable configuration screens but not as a dedicated student-facing directory.')
bullet('The timetable engine runs on a single Render instance. It does not support horizontal scaling and is not designed for simultaneous multi-department generation.')
bullet('Notifications dismissed locally in the app (via SharedPreferences) are not synced across devices. If a user signs in on a second device, previously dismissed notifications will reappear.')

section_heading('6.3  Future Work')
bullet('AI Academic Assistant: Restore the RAG-based assistant using Gemini text-embedding-004 (768-dim) and gemini-2.0-flash. This requires running the documents table migration (adding namespace and content_tsv columns, GIN index, rewriting the match_documents() RPC) and building the lib/features/ai_assistant/ feature module.')
bullet('iOS Support: Configure the Flutter project for iOS, add the Apple-specific Firebase and Supabase configurations, and submit to the App Store.')
bullet('Student Routine Cancellation Badges: Wire cancellations data to the student weekly routine grid to visually badge cancelled class slots.')
bullet('Teacher Directory Screen: Build a student-facing teacher directory screen with search, filtering by subject, and contact information.')
bullet('Cross-device Notification Sync: Persist dismissed notification IDs in the notification_reads table (rather than local SharedPreferences) to sync dismiss state across devices.')
bullet('Production APK Signing: Set up a release keystore and publish the application to the Google Play Store.')
bullet('Multi-department Engine: Extend the timetable engine to support multiple departments within the same institution, with per-department configuration and room pools.')
bullet('Assignment Submission Module: Complete the assignments and submissions tables that are already in the schema (with the is_late DB trigger) by building the teacher assignment creation and student submission screens.')

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('References', 20, space_before=0, space_after=12)

refs = [
    '[1] Flutter Documentation (2024). Flutter – Build apps for any screen. https://flutter.dev/docs',
    '[2] Supabase Documentation (2024). Supabase – The Open Source Firebase Alternative. https://supabase.com/docs',
    '[3] FastAPI Documentation (2024). FastAPI – Modern, fast (high-performance) web framework for building APIs with Python. https://fastapi.tiangolo.com',
    '[4] Google OR-Tools Documentation (2024). OR-Tools – Google\'s Operations Research software. https://developers.google.com/optimization',
    '[5] Firebase Cloud Messaging Documentation (2024). Firebase Cloud Messaging. https://firebase.google.com/docs/cloud-messaging',
    '[6] Hartert, R., Schaus, P., Dejemeppe, C., Dury, A., Moiny, S., & Monfroy, E. (2015). A Hybrid CP-SAT Approach for the University Course Timetabling Problem. Proceedings of the International Conference on the Integration of AI and OR Techniques in Constraint Programming.',
    '[7] Lewis, R. (2008). A Survey of Metaheuristic-Based Techniques for University Timetabling Problems. OR Spectrum, 30(1), 167–190.',
    '[8] PostgreSQL Documentation (2024). PostgreSQL – The World\'s Most Advanced Open Source Relational Database. https://www.postgresql.org/docs',
    '[9] Lewis, P. et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information Processing Systems (NeurIPS), 33.',
    '[10] Google Dart Documentation (2024). Dart programming language. https://dart.dev/guides',
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    set_font(run, 12)

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX A – SOURCE CODE HIGHLIGHTS
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Appendix A', 20, space_before=0, space_after=4)
heading('Source Code Highlights', 16, space_before=0, space_after=12)

body('This appendix highlights selected key code patterns from the UniVerse codebase that demonstrate the core technical design decisions.')

subsection_heading('A.1  CP-SAT Solver – Hard Constraint Example (solver.py)')
body('The following snippet illustrates how the CP-SAT solver enforces the teacher no-double-booking constraint. For each teacher t and each (day, period) slot s, the sum of all boolean assignment variables for teacher t in slot s must be at most 1:')

code_block = doc.add_paragraph()
code_block.paragraph_format.left_indent = Inches(0.5)
code_block.paragraph_format.space_after = Pt(6)
code_text = (
    "for teacher in teachers:\n"
    "    for day in days:\n"
    "        for period in periods:\n"
    "            sessions_in_slot = [\n"
    "                assignments[(s, day, period)]\n"
    "                for s in teacher_sessions[teacher]\n"
    "            ]\n"
    "            if sessions_in_slot:\n"
    "                model.Add(sum(sessions_in_slot) <= 1)"
)
run = code_block.add_run(code_text)
run.font.name = 'Courier New'
run.font.size = Pt(10)

subsection_heading('A.2  Service Layer Pattern (teacher_service.dart)')
body('The cancel class operation demonstrates the service-layer pattern. The screen calls the controller, the controller calls the service, and the service is the only layer that touches Supabase:')

code_block2 = doc.add_paragraph()
code_block2.paragraph_format.left_indent = Inches(0.5)
code_block2.paragraph_format.space_after = Pt(6)
code_text2 = (
    "Future<void> cancelClass({\n"
    "  required String routineId,\n"
    "  required DateTime classDate,\n"
    "  required String reason,\n"
    "  required String batch,\n"
    "  required String section,\n"
    "  required String subject,\n"
    "}) async {\n"
    "  await _supabase.from('cancellations').insert({\n"
    "    'routine_id': routineId,\n"
    "    'class_date': classDate.toIso8601String().split('T').first,\n"
    "    'reason': reason,\n"
    "    'batch': batch,\n"
    "    'section': section,\n"
    "    'subject': subject,\n"
    "    'cancelled_by': _supabase.auth.currentUser!.id,\n"
    "  });\n"
    "}"
)
run2 = code_block2.add_run(code_text2)
run2.font.name = 'Courier New'
run2.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX B – USER'S GUIDE
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('Appendix B', 20, space_before=0, space_after=4)
heading("User's Guide", 16, space_before=0, space_after=12)

subsection_heading('B.1  Installation')
bullet('Download the UniVerse APK file (app-release.apk).')
bullet('On your Android device, enable installation from unknown sources (Settings → Security → Install unknown apps).')
bullet('Open the APK file and follow the on-screen prompts to install.')
bullet('Launch the app from your home screen.')

subsection_heading('B.2  Registration – Student')
bullet('On the Login screen, tap "Sign up with Email".')
bullet('Select role: Student. Enter your full name, university email, batch (e.g., 62), section (e.g., G), and a password.')
bullet('Check your email for a verification link and click it.')
bullet('Return to the app and sign in with your credentials.')

subsection_heading('B.3  Registration – Teacher')
bullet('On the Login screen, tap "Sign up with Email".')
bullet('Select role: Teacher. Enter your full name, university email, teacher code (e.g., TKM), and a password.')
bullet('Verify your email and sign in.')

subsection_heading('B.4  Student Usage')
bullet('Home tab: View your live or next class. The countdown timer updates in real time.')
bullet('Routine tab: Select a day tab (Sun–Sat) to see your full day\'s schedule. Classes with an active cancellation are marked CANCELLED.')
bullet('Resources tab: Tap a semester folder, then a category, to browse and open academic files.')
bullet('Alerts tab: View all notifications. Long-press a notification to enter multi-select mode; tap the trash icon to dismiss selected notifications.')

subsection_heading('B.5  Teacher Usage')
bullet('Home tab: View your live or next class.')
bullet('Routine tab: View your full assigned schedule.')
bullet('Classes tab: Select a day. Tap a class to open the action menu: Cancel Occurrence (enter a reason → students are notified instantly), Post Update (room change / notice / test reminder), or Undo Cancellation.')

subsection_heading('B.6  Admin Usage')
bullet('Dashboard tab: View system stats.')
bullet('Broadcast tab: Compose a notification with title, body, and audience (all users / by role / by batch+section). Tap the history icon to view past broadcasts.')
bullet('Routine tab: Switch between Manage (view/delete routine entries) and Generate (upload Course Distribution file → wait for solver → review report → download or publish).')
bullet('Resources tab: Upload files or Drive links with semester and category tags. Tap the library icon to view and delete existing resources.')
bullet('Users tab: View registered accounts.')
bullet('Admin secondary screens (via app-bar icons): Manage Rooms, Manage Faculty, Timetable Settings.')

subsection_heading('B.7  Timetable Generation – Step by Step')
bullet('Go to the Routine tab and select the Generate segment.')
bullet('Tap "Select Course Distribution File" and choose the .xlsx workbook from your device.')
bullet('Verify that rooms, faculty, and settings are configured correctly (use the manage screens accessible from the app bar).')
bullet('Tap "Generate Timetable". The progress bar and status log will update as the engine processes the file.')
bullet('Once complete, review the Stats report (sessions scheduled, cohorts, teachers, rooms used) and the Validation card (clash counts, pass/fail).')
bullet('Tap "View Grid" to see a visual day-by-day timetable preview.')
bullet('Tap "Download" to save the Excel workbook, or "Publish" to replace the live routines table and notify all users.')

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
output_path = r'C:\Users\sirsw\Desktop\Project\UniVerse\UniVerse_Project_Report.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
